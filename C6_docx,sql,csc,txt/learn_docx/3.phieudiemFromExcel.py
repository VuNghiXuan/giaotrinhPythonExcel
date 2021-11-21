from docx import Document
from docx.shared import Inches
import xlwings as xw

def tim_VungChuaDuLieu(sh):
    # Sử dụng hàm có sẵn last_cell.row để tìm dòng cuối cùng của bảng tính
    lr_table = sh.cells.last_cell.row
    # print('Dòng cuối cùng bảng tính là:', lr_table)

    # Sử dụng hàm có sẵn last_cell.column để tìm cột cuối cùng của bảng tính
    lc_table = sh.cells.last_cell.column
    # print('Cột cuối cùng bảng tính là:', lc_table)

    " Thêm code để tìm dòng cuối, cột cuối có chứa dữ liệu";
    lr_data = sh.range('A'+ str(lr_table)).end('up').row
    # print('Dòng cuối cùng có chứa dữ liệu là:', lr_data)

    # lcol = sh.range(row_index, col).end("left").column
    lc_data = sh.range(1, lc_table).end('left').column
    # print('Cột cuối cùng có chứa dữ liệu là:', lc_data)

    # Cuối cùng ta có bảng dữ liệu cần thu thập như sau:
    table_datas = sh.range((1,1), (lr_data,lc_data))
    # print("Bảng dữ liệu có địa chỉ là:", table_datas.address)
    return table_datas


"Phần input ";
# file Excel
pathExcel = "D:\\MyBook_PyExcel\\learn_docx\\1.BangDiem.xlsx"

# Sử dụng cặ từ khóa để bỏ qua lỗi lúc chưa tắt file Excel
try: 
    excel = xw.Book(pathExcel)
except:
    print(f"file Excel {pathExcel} đang mở")
sh = excel.sheets.active
datasFromExcel = tim_VungChuaDuLieu(sh)
datasFromExcel = datasFromExcel.value
# datasFromExcel = excel.sheets('DS_xeploai').range('A1:J6').value

# file hình
images = ['1.png','2.png','3.png','4.png','5.png','6.png','7.png','8.png']

# Phân tiêu đề và data vào 2 list
title_datas = datasFromExcel[0]

for i_r in range(1, len(datasFromExcel)):
    
    # Tạo file Word và lấy thông tin do thí sinh cung cấp    
    doc = Document()
    # Tiêu đề đầu tiên Word 
    doc.add_heading('PHIẾU BÁO ĐIỂM', level = 0) # level = 0 : tham số cỡ chữ lớn nhất trong các tiêu đề

    # Dữ liệu thí sinh tự điền: khóa học, tên, ngày sinh
    thongtinThiSinhCungCap = datasFromExcel[i_r][0:4]

    # Tên khóa học
    tieude_KhoaHoc = title_datas[0]
    ghiTieuDeKhoaHoc = doc.add_paragraph(f"{tieude_KhoaHoc}: ")
    tenKhoaHoc = thongtinThiSinhCungCap[0]    
    ghiTieuDeKhoaHoc.add_run (f'{tenKhoaHoc}').bold = True

    # Thông tin thí sinh:
    doc.add_heading('Thông tin thí sinh', level = 1)

    # Chèn hình học sinh: 
    try:
        image = doc.add_picture(f'./image/{images[i_r-1]}', width = Inches(1.25))
    except: 
        print("Không có ảnh")

    # Số báo danh
    tieudeSoBaoDanh = title_datas[1]
    ghiSoBaoDanh = doc.add_paragraph(f"{tieudeSoBaoDanh}: ", style = "List Bullet")
    soBaoDanh = str(int(thongtinThiSinhCungCap[1]))     
    ghiSoBaoDanh.add_run (f'{soBaoDanh}').bold = True

    # Tên thí sinh
    tieudeTenThiSinh = title_datas[2]
    ghiTenThiSinh = doc.add_paragraph(f"{tieudeTenThiSinh}: ", style = "List Bullet")
    tenThiSinh = str(thongtinThiSinhCungCap[2])     
    ghiTenThiSinh.add_run (f'{tenThiSinh}').bold = True

    # Ngày sinh
    tieudeNgaySinh = title_datas[3]
    ghiNgaySinh = doc.add_paragraph(f"{tieudeNgaySinh}: ", style = "List Bullet")
    ngaySinh = (thongtinThiSinhCungCap[3]).strftime("%d/%m/%Y")     
    ghiNgaySinh.add_run (f'{ngaySinh}').bold = True

    # Thêm nội dung điểm
    doc.add_heading('Thông tin điểm thi', level = 1)


    # list điểm
    listDiem = datasFromExcel[i_r][4:]
    for i_diem in range(len(listDiem)):        
        # Nhập điểm
        tieudeDiem = title_datas[4+i_diem]
        ghiDiem = doc.add_paragraph(f"{tieudeDiem}: ", style = "List Number")
        
        diem = datasFromExcel[i_r][4+i_diem]
        try: 
            diem = str(round(float(diem),1))+ " (điểm)"
        except:
            pass
        diem = str(diem)     
        ghiDiem.add_run (f'{diem}').bold = True

        # Lưu file word 
    doc.save(f'./out_Word/{i_r}.{tenThiSinh}.docx')
        
        