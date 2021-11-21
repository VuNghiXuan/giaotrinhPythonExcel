from docx import Document
from docx.shared import Inches
# Gán cái Document 
doc = Document()
# Add tiêu đề đầu tiên là 
doc.add_heading('Phiếu báo điểm', level = 0) # level = 0 : tham số cỡ chữ lớn nhất trong các tiêu đề

"Thêm các nội dung cho phiếu ------"
# Thêm nội dung
content = doc.add_paragraph("Tên thí sinh: ")
# Thêm chú thích cho nội dung và in đậm
ten = content.add_run("Vũ Nghi Xuân").bold = True
lop = content.add_run (' - Lớp: Python Excel').italic = True

# Chèn hình học sinh: 
image = doc.add_picture('./picture/anhVu.png', width = Inches(1.25))

# Thêm nội dung điểm
doc.add_heading('Điểm các môn', level = 1)

# Nhập điểm toán
toan = doc.add_paragraph("Toán: ", style = "List Bullet")
toan.add_run('8 điểm')
# Nhập điểm lý
ly = doc.add_paragraph("Lý: ", style = "List Bullet")
ly.add_run('5 điểm')
# Nhập điểm hóa
hoa = doc.add_paragraph("Hóa: ", style = "List Bullet")
hoa.add_run('10 điểm')

# Thêm nội dung Xếp loại
doc.add_heading('Xếp loại', level = 1)
# Nhập Xếp loại
xeploai = doc.add_paragraph("Xếp loại: ", style = "List Number")
xeploai.add_run('khá')
# Nhập Xếp loại
danhgia = doc.add_paragraph("Đánh giá: ", style = "List Number")
danhgia.add_run('Được lên lớp')

# Lưu file word 
doc.save('phieudiem.docx')